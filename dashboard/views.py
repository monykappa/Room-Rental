from django.shortcuts import render
from django.shortcuts import render, redirect
from django.shortcuts import render, get_object_or_404
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt  
from django.contrib.auth import login, authenticate
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from django.contrib.auth.hashers import make_password
from django.contrib.auth.models import User 
from django.contrib.auth import logout
from .models import *
from django.utils import timezone
from django.contrib.auth.views import LogoutView
from django.db.models import Q
from decimal import Decimal
from django.views.generic import ListView, DeleteView, UpdateView
from django.urls import reverse_lazy
from django.views.decorators.http import require_POST
from .forms import *
from .models import CheckIn as CheckInModel
from django.urls import reverse
from django.db.models import Max
from datetime import timedelta
from django.conf import settings
from django.http import HttpResponse
from django.core.exceptions import MultipleObjectsReturned, ObjectDoesNotExist
from django.views import View
from django.utils.translation import gettext_lazy as _
from django.contrib.auth.mixins import LoginRequiredMixin
from django.forms import formset_factory
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.forms.models import modelformset_factory
from django.http import HttpResponseBadRequest
from django.db import transaction 
from datetime import datetime
from reportlab.lib.pagesizes import landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, PageTemplate, BaseDocTemplate, PageBreak, Paragraph 
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
import os





def signin_view(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)

        if user is not None:
            login(request, user)
            # Redirect to a success page or any other page after login
            return redirect('dashboard:dashboard')
        else:
            # Handle invalid login credentials
            # You may customize this part based on your requirements
            return render(request, 'userprofile/sign_in.html', {'error_message': 'Invalid username or password'})

    return render(request, 'userprofile/sign_in.html')


def logout_view(request):
    logout(request)
    return redirect('dashboard:signin') 

@login_required
def base(request):
    template = 'base.html'  
    return render(request, template)

@login_required
def dashboard(request):
    rooms = room.objects.all()
    template = 'dashboard/dashboard.html'
    return render(request, template, {'rooms': rooms})

@login_required
def add_room(request):
    if request.method == 'POST':
        form = RoomForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect(reverse('dashboard:room_list'))  # Use the correct name here
        else:
            # Handle form errors if needed
            pass
    else:
        form = RoomForm()

    house_owners = HouseOwner.objects.all()
    return render(request, 'dashboard/room/add_room.html', {'house_owners': house_owners, 'form': form})

@login_required
def room_list(request):
    rooms = room.objects.annotate(latest_checkin=Max('check_in_entries__date')).order_by('-latest_checkin')
    house_owners = HouseOwner.objects.all()
    return render(request, 'dashboard/room/room_list.html', {'rooms': rooms, 'house_owners': house_owners})


class RoomUpdateView(UpdateView, LoginRequiredMixin, View):
    login_url = '/dashboard:signin_view/'
    model = room
    template_name = 'dashboard/room/edit_room.html'
    form_class = RoomForm  
    checkin_form_class = CheckInForm


    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        try:
            checkin_instance = CheckIn.objects.get(room=self.get_object())
            context['checkin_form'] = self.checkin_form_class(instance=checkin_instance)
        except CheckIn.DoesNotExist:
            # Handle the case when CheckIn does not exist for the room
            pass
        return context

    def form_valid(self, form):
        checkin_instance = None
        try:
            checkin_instance = CheckIn.objects.get(room=self.get_object())
        except CheckIn.DoesNotExist:
            # Handle the case when CheckIn does not exist for the room
            pass

        if checkin_instance:
            checkin_form = self.checkin_form_class(self.request.POST, instance=checkin_instance)
            if checkin_form.is_valid():
                checkin_form.save()

        return super().form_valid(form)

    
    labels = {
        'RoomNo': _('Room No / លេខ​បន្ទប់'),
        'HouseOwner': _('House Owner / ម្ចាស់ផ្ទះ'),
        'RoomFee': _('Room Price / តម្លៃបន្ទប់ $'),
        'remark': _('Remark / ចំណាំ'),
        'status': _(' Status / ស្ថានភាព'),
    }
    placeholders = {
        'RoomNo': _('បញ្ចូលលេខ​បន្ទប់'),
        'HouseOwner': _('បញ្ចូលឈ្មោះម្ចាស់ផ្ទះ'),
        'RoomFee': _('បញ្ចូលតម្លៃបន្ទប់ $'),
        'remark': _('បញ្ចូលចំណាំ'),
        'status': _('ជ្រើសរើសស្ថានភាព'),
    }

    def get_form(self, form_class=None):
        form = super(RoomUpdateView, self).get_form(form_class)

        # Try to retrieve the related Checkin instance
        try:
            checkin_instance = CheckIn.objects.get(room=self.get_object())
            client_name = checkin_instance.client.ClientName if checkin_instance and checkin_instance.client else ""
        except CheckIn.DoesNotExist:
            client_name = ""

        form.fields['client_name'].initial = client_name
        form.fields['client_name'].widget.attrs['value'] = client_name

        return form

    success_url = reverse_lazy('dashboard:room_list')

class RoomDeleteView(DeleteView, LoginRequiredMixin, View):
    login_url = '/dashboard:signin_view/'
    model = room
    template_name = 'dashboard/room/delete_room.html'
    success_url = reverse_lazy('dashboard:room_list') 

@login_required
def house_owner(request):
    house_owners = HouseOwner.objects.all().order_by('-id')
    return render(request, 'dashboard/house_owner/house_owner.html', {'house_owners': house_owners})

def add_house_owner(request):
    if request.method == 'POST':
        form = HouseOwnerForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('dashboard:house_owner')  # Redirect to your dashboard
    else:
        form = HouseOwnerForm()

    return render(request, 'dashboard/house_owner/add_house_owner.html', {'form': form})

class HouseOwnerEditView(View, LoginRequiredMixin):
    login_url = '/dashboard:signin_view/'
    def get(self, request, pk):
        house_owner = get_object_or_404(HouseOwner, pk=pk)
        form = HouseOwnerForm(instance=house_owner)
        return render(request, 'dashboard/house_owner/house_owner_edit.html', {'form': form, 'house_owner': house_owner})

    def post(self, request, pk):
        house_owner = get_object_or_404(HouseOwner, pk=pk)
        form = HouseOwnerForm(request.POST, instance=house_owner)
        if form.is_valid():
            form.save()
            return redirect('dashboard:house_owner')  # Adjust the URL based on your project
        return render(request, 'dashboard/house_owner/house_owner_edit.html', {'form': form, 'house_owner': house_owner})
    
    def delete(self, request, pk):
        house_owner = get_object_or_404(HouseOwner, pk=pk)
        house_owner.delete()
        return JsonResponse({'message': 'House owner deleted successfully.'})
    

class HouseOwnerDeleteView(View, LoginRequiredMixin):
    login_url = '/dashboard:signin_view/'
    def get(self, request, pk):
        house_owner = get_object_or_404(HouseOwner, pk=pk)
        return render(request, 'delete_template.html', {'house_owner': house_owner})

    def post(self, request, pk):
        house_owner = get_object_or_404(HouseOwner, pk=pk)
        house_owner.delete()
        return redirect('house_owner_list')  # Redirect to your list view


@login_required
def client(request):
    clients = Client.objects.all().order_by('-id')
    return render(request, 'dashboard/client.html', {'clients': clients})



class ClientListView(ListView, LoginRequiredMixin, View):
    login_url = '/dashboard:signin_view/'
    model = Client
    template_name = 'dashboard/client.html'
    context_object_name = 'clients'

class ClientEditView(UpdateView, LoginRequiredMixin, View):
    login_url = '/dashboard:signin_view/'
    model = Client
    form_class = ClientForm
    template_name = 'dashboard/client_edit_form.html'
    success_url = reverse_lazy('dashboard:client')


    def form_valid(self, form):
        # Retrieve the client object
        client = get_object_or_404(Client, pk=self.kwargs['pk'])

        # Update the client name in the associated CheckIn entries
        client.update_checkin_entries()

        return super().form_valid(form)
    
    def get(self, request, *args, **kwargs):
        self.object = self.get_object()
        context = self.get_context_data(object=self.object, form=self.get_form())
        return self.render_to_response(context)

    def post(self, request, *args, **kwargs):
        self.object = self.get_object()
        form = self.get_form()
        if form.is_valid():
            return self.form_valid(form)
        else:
            return self.form_invalid(form)


class ClientDeleteView(DeleteView, LoginRequiredMixin, View):
    login_url = '/dashboard:signin_view/'
    model = Client
    template_name = 'dashboard/client_confirm_delete.html'
    success_url = reverse_lazy('dashboard:client-list-class-based')


@login_required
def check_in(request):
    check_ins = CheckInModel.objects.all().order_by('-id') 
    return render(request, 'dashboard/check_in.html', {'CheckIn': check_ins})


@login_required
def checkin_form(request):
    if request.method == 'POST':
        form = CheckInForm(request.POST)
        if form.is_valid():
            # Create or get the Client instance
            client_instance, created = Client.objects.get_or_create(
                ClientName=form.cleaned_data['client_name'],
                address=form.cleaned_data['client_address'],
                contact=form.cleaned_data['client_contact']
            )

            # Update the form's fields with the Client instance
            form.instance.client_name = client_instance.ClientName
            form.instance.client_address = client_instance.address
            form.instance.client_contact = client_instance.contact

            # Save the Client instance
            client_instance.save()

            # Save the CheckIn form data
            form.save()

            # Redirect to the check-in list page or any other page
            return redirect('dashboard:CheckIn')

    else:
        form = CheckInForm()

    return render(request, 'dashboard/checkin_form.html', {'form': form})

@login_required
def checkout_view(request, checkin_id):
    checkin = get_object_or_404(CheckIn, pk=checkin_id)

    # Create a CheckOut entry
    CheckOut.objects.create(ClientName=checkin.client, room=checkin.room)

    # Update room status
    checkin.room.status = 'Available/ទំនេរ'
    checkin.room.save()

    return redirect('dashboard:check_in') 


@login_required
def check_out(request):
    check_outs = CheckOut.objects.all().order_by('-id') 

    for checkout in check_outs:
        try:
            checkin = CheckIn.objects.get(client_name=checkout.ClientName.ClientName)
            checkout.total_stay = (checkout.date - checkin.date).days
            checkout.checkin_date = checkin.date
        except CheckIn.DoesNotExist:
            checkout.total_stay = None
            checkout.checkin_date = None

    return render(request, 'dashboard/check_out.html', {'CheckOut': check_outs})



@login_required
def checkout_form_view(request):
    if request.method == 'POST':
        form = check_out_form(request.POST)
        if form.is_valid():
            checkout_instance = form.save(commit=False)
            
            # Update the room status and save the checkout entry
            checkout_instance.room.status = 'Available/ទំនេរ'
            checkout_instance.room.save()
            
            # Automatically get the associated client of the selected room during check-in
            checkin_entry = checkout_instance.room.check_in_entries.last()
            
            if checkin_entry and hasattr(checkin_entry, 'client_name'):
                client_name = checkin_entry.client_name
                # Assuming 'Client' has a field 'ClientName', modify if needed
                client_instance, created = Client.objects.get_or_create(ClientName=client_name)
                checkout_instance.ClientName = client_instance

            checkout_instance.save()
            return redirect('dashboard:CheckOut')
    else:
        form = check_out_form()

    return render(request, 'dashboard/checkout_form.html', {'form': form})


@login_required
def other_fee(request):
    trashes = Trash.objects.all()
    parkings = parking.objects.all()

    return render(request, 'dashboard/other_fee.html', {'trashes': trashes, 'parkings': parkings})



@login_required
def edit_trash(request, trash_id):
    trash_instance = get_object_or_404(Trash, id=trash_id)
    success_message = None

    if request.method == 'POST':
        new_price = request.POST.get('new_price')
        trash_instance.TrashPrice = new_price
        trash_instance.save()
        success_message = "Trash price updated successfully"

        # Redirect to the same page to fetch the updated data
        return redirect('dashboard:other_fee')

    return render(request, 'dashboard/other_fee.html', {'success_message': success_message})



@login_required
def edit_parking(request, parking_id):
    parking_instance = get_object_or_404(parking, id=parking_id)
    success_message = None

    if request.method == 'POST':
        new_price = request.POST.get('new_price')
        parking_instance.ParkingPrice = new_price
        parking_instance.save()
        success_message = "Parking price updated successfully"

        # Redirect to the same page to fetch the updated data
        return redirect('dashboard:other_fee')

    return render(request, 'dashboard/other_fee.html', {'success_message': success_message})



@login_required
def monthly_fee(request):
    monthly_fees = MonthlyRentalFee.objects.all().order_by('-id')

    # Process the month filter form
    initial_month = datetime.now().month
    initial_year = datetime.now().year
    form = MonthFilterForm(initial={'selected_month': initial_month, 'selected_year': initial_year})

    if request.GET.get('selected_month') and request.GET.get('selected_year'):
        form = MonthFilterForm(request.GET)

        if form.is_valid():
            selected_month = int(form.cleaned_data['selected_month']) if form.cleaned_data['selected_month'] else None
            selected_year = int(form.cleaned_data['selected_year']) if form.cleaned_data['selected_year'] else None
            if selected_month is not None and selected_year is not None:
                monthly_fees = monthly_fees.filter(date__month=selected_month, date__year=selected_year)

    # Set the number of items per page
    items_per_page = 10  # You can adjust this value

    paginator = Paginator(monthly_fees, items_per_page)
    page = request.GET.get('page')

    try:
        monthly_fees = paginator.page(page)
    except PageNotAnInteger:
        # If page is not an integer, deliver the first page.
        monthly_fees = paginator.page(1)
    except EmptyPage:
        # If page is out of range (e.g. 9999), deliver the last page of results.
        monthly_fees = paginator.page(paginator.num_pages)

    return render(request, 'dashboard/monthly/monthly_rental_fee.html', {'monthly_fees': monthly_fees, 'form': form})


@login_required
def add_monthly_fee(request):
    if request.method == 'POST':
        forms = [MonthlyRentalFeeForm(room, request.POST, prefix=f'form-{room.id}') for room in room.objects.all()]
        if all(form.is_valid() for form in forms):
            for form in forms:
                form.save()
            return redirect('dashboard:monthly_fee')

    forms = [MonthlyRentalFeeForm(room, prefix=f'form-{room.id}') for room in room.objects.all()]
    previous_waters = [form.get_current_water() for form in forms]
    context = {'forms': zip(forms, previous_waters)}
    return render(request, 'dashboard/monthly/add_monthly_fee.html', context)





from reportlab.pdfgen import canvas
from io import BytesIO
from reportlab.lib.pagesizes import letter


from reportlab.lib import colors
from django.template.loader import get_template
from xhtml2pdf import pisa
import locale



@login_required
def export_to_pdf(request, selected_month, selected_year):
    font_path = os.path.join(settings.BASE_DIR, 'fonts', 'KhmerOS_siemreap.ttf')
    locale.setlocale(locale.LC_ALL, '')

    # Register the Khmer font with the "Middle Eastern and South Asian" text engine
    pdfmetrics.registerFont(TTFont('Khmer OS Siemreap', font_path))
    pdfmetrics.registerFont(TTFont('Khmer OS Muol', os.path.join(settings.BASE_DIR, 'fonts', 'KhmerOS_muol.ttf')))
    pdfmetrics.registerFont(TTFont('Khmer OS Bokor', os.path.join(settings.BASE_DIR, 'fonts', 'KhmerOS_bokor.ttf')))
    pdfmetrics.registerFont(TTFont('Khmer OS Siemreap', os.path.join(settings.BASE_DIR, 'fonts', 'KhmerOS_siemreap.ttf')))

    # Retrieve the monthly fees based on the selected month and year
    monthly_fees = MonthlyRentalFee.objects.filter(date__month=selected_month, date__year=selected_year)

    # Create a file-like buffer to receive PDF data
    buffer = BytesIO()

    # Create the PDF object using the buffer as its "file"
    p = SimpleDocTemplate(buffer, pagesize=landscape(letter))

    # Set up the PDF document
    elements = []

    for i, monthly_fee in enumerate(monthly_fees):
        # Add a new page for each room except the first one
        if i > 0:
            elements.append(PageBreak())

        # Add the title for each new page
        title_text = f"Monthly rental for / ការជួលប្រចាំខែ{selected_month}/{selected_year}"
        title = Paragraph(title_text, getSampleStyleSheet()['Title'])
        elements.append(title)

        # Set up the PDF document for the new page
        data = []
        table_headers = [
            "ល.រ",
            "ម្ចាស់ផ្ទះ",
            "កាលបរិច្ឆេទ",
            "ទឹកប្រាក់បន្ថែម",
            "ទឹកប្រាក់ថ្មី",
            "តម្លៃទឹក",
            "តម្លៃបន្ទប់",
            "តម្លៃពោធិ",
            "តម្លៃចល័ត",
            "សរុប($)",
            "សរុប(៛)",
        ]
        data.append(table_headers)

        # Format the total in Khmer Riel using locale
        formatted_riel_total = locale.format_string('%.0f', monthly_fee.total * 4100, grouping=True)

        data_row = [
            str(monthly_fee.room.RoomNo),
            str(monthly_fee.room.HouseOwner.name),
            str(monthly_fee.date),
            f"{monthly_fee.previous_water} m³",
            f"{monthly_fee.current_water} m³",
            f"${monthly_fee.water_fee:.2f}",
            f"${monthly_fee.room.RoomFee:.2f}",
            f"${monthly_fee.trash_fee:.2f}",
            f"${monthly_fee.park_fee:.2f}",
            f"${monthly_fee.total:.2f}",
            formatted_riel_total,
        ]
        data.append(data_row)

        # Create the table and set style
        table = Table(data)
        style = TableStyle([('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Khmer OS Siemreap'),  # Set the font for the table
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black)])

        table.setStyle(style)

        # Add the table to the elements for this page
        elements.append(table)

    # Build the PDF
    p.build(elements)

    # FileResponse sets the Content-Disposition header so that browsers
    # present the option to save the file.
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'filename="monthly_rental_fee_report_{selected_month}_{selected_year}.pdf"'
    return response


import openpyxl
from openpyxl import Workbook
from openpyxl.styles import Font
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

from docx.oxml import OxmlElement
locale.setlocale(locale.LC_ALL, 'en_US.UTF-8')



@login_required
def export_to_excel(request, selected_month, selected_year):
    # Retrieve the monthly fees based on the selected month and year
    monthly_fees = MonthlyRentalFee.objects.filter(date__month=selected_month, date__year=selected_year)

    # Create a new Excel workbook and add a worksheet
    workbook = Workbook()
    worksheet = workbook.active

    # Add headers to the worksheet
    headers = ["លេខ​បន្ទប់", "ម្ចាស់ផ្ទះ", "កាលបរិច្ឆេទ", "ទឹកខែមុន", "ទឹកបច្ចុប្បន្ន", "តម្លៃទឹក", "តម្លៃបន្ទប់", "ថ្លៃសំរាម", "ថ្លៃចតរថយន្ត", "សរុប($)", "សរុប(៛)"]
    worksheet.append(headers)

    # Add data to the worksheet
    for monthly_fee in monthly_fees:
        row_data = [
            monthly_fee.room.RoomNo,
            monthly_fee.room.HouseOwner.name,
            str(monthly_fee.date),
            monthly_fee.previous_water,
            monthly_fee.current_water,
            monthly_fee.water_fee,
            monthly_fee.room.RoomFee,
            monthly_fee.trash_fee,
            monthly_fee.park_fee,
            monthly_fee.total,
            monthly_fee.total * 4100,  # Assuming this is the conversion to Khmer Riel
        ]
        worksheet.append(row_data)

    # Create a response with the Excel file
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename=monthly_rental_fee_report_{selected_month}_{selected_year}.xlsx'
    workbook.save(response)

    return response



from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, RGBColor
import calendar

def set_landscape_section(section):
    section_start = section.start_type
    section.orientation = WD_ORIENT.LANDSCAPE
    section.start_type = section_start


def add_border(cell):
    if cell.paragraphs and cell.paragraphs[0].runs:
        # Customize border
        border = cell.paragraphs[0].runs[0].font
        border.size = Pt(1)  # Border width
        # Add padding
        cell.paragraphs[0].paragraph_format.left_indent = Pt(4)
        cell.paragraphs[0].paragraph_format.right_indent = Pt(4)

    
def export_to_word(request, selected_month, selected_year):
    # Create a new Word document
    document = Document()
    # Retrieve monthly rental fees based on the selected month and year
    monthly_fees = MonthlyRentalFee.objects.filter(date__month=selected_month, date__year=selected_year)

    # Loop through each monthly fee entry
    for i, monthly_fee in enumerate(monthly_fees):
        # Add a new section (page) for each room except the first one
        if i > 0:
            section = document.add_section()
            set_landscape_section(section)

        # Add the title for each new page
        room_number = monthly_fee.room.RoomNo

# Define a mapping of English month names to Khmer month names
        khmer_month_names = {
            'January': 'មករា',
            'February': 'កុម្ភៈ',
            'March': 'មីនា',
            'April': 'មេសា',
            'May': 'ឧសភា',
            'June': 'មិថុនា',
            'July': 'កក្កដា',
            'August': 'សីហា',
            'September': 'កញ្ញា',
            'October': 'តុលា',
            'November': 'វិច្ឆិកា',
            'December': 'ធ្នូ'
        }

        english_month_name = calendar.month_name[selected_month]
        khmer_month_name = khmer_month_names.get(english_month_name, 'Unknown')

        formatted_date = f"{khmer_month_name}-{selected_year}"

        title_text = f"ថ្លៃជួលប្រចាំខែ \nបន្ទប់លេខ {room_number}\n{monthly_fee.room.HouseOwner.name}\n{formatted_date}"
        title = document.add_paragraph()
        title_run = title.add_run(title_text)
        title_run.bold = True
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run.font.size = Pt(18)  # Increase font size for the title

        # Set Khmer OS font for the title text
        title_run.font.name = 'Khmer OS'


        # Create a table for each room with 2 columns and 10 rows
        table = document.add_table(rows=10, cols=2)
        table.alignment = WD_ALIGN_VERTICAL.CENTER  # Center-align content vertically

        # Set table style to include borders
        table.style = 'Table Grid'

        # Customize borders for each cell in the table
        for row in table.rows:
            for cell in row.cells:
                add_border(cell)
                
            utilities = Utilities.objects.first()  
            water_rate = utilities.water_rate if utilities else 0
        # Add content to the table
            content_labels = ["លេខ​បន្ទប់", "កាលបរិច្ឆេទ", "ទឹកខែមុន", "ទឹកបច្ចុប្បន្ន",
                            "អត្រាទឹក", "ថ្លៃទឹក", "ថ្លៃ​ចតរថយន្ត", "ថ្លៃសំរាម", "ថ្លៃបន្ទប់", "ថ្លៃសរុប"]

            # Calculate total in Riel using an exchange rate (1 USD = 4000 Riel)
            exchange_rate = 4100
            total_riel = monthly_fee.total * exchange_rate
            
            formatted_total_riel = "{:,.0f}".format(total_riel) if total_riel % 1 == 0 else "{:,.2f}".format(total_riel)
            if water_rate is not None:
                formatted_water_rate = "{:.2f}".format(water_rate.rate) if water_rate is not None else "N/A"

            else:
                formatted_water_rate = "N/A"  # Or any default value you want to use when water_rate is None


            content_values = [
                str(monthly_fee.room.RoomNo),
                monthly_fee.date.strftime('%d-%m-%Y'),
                f"{monthly_fee.previous_water} m³",
                f"{monthly_fee.current_water} m³",
                f"${formatted_water_rate}/m³",  # Include "/m³" in the water rate
                "${:,.0f}".format(monthly_fee.water_fee),
                "${:,.0f}".format(monthly_fee.park_fee),
                "${:,.0f}".format(monthly_fee.trash_fee),
                "${:,.0f}".format(monthly_fee.room.RoomFee),
                "${:,.0f} រឺ {:,.0f} ៛".format(monthly_fee.total, total_riel),
    ]

        # Iterate through labels and values to fill the table
        for row_index, (label, value) in enumerate(zip(content_labels, content_values)):
            table.cell(row_index, 0).text = label
            table.cell(row_index, 1).text = value
            
            # Adjust cell margins for padding
            for cell in [table.cell(row_index, 0), table.cell(row_index, 1)]:
                cell.paragraphs[0].runs[0].font.size = Pt(14)
                cell.paragraphs[0].runs[0].font.name = 'Khmer OS Battambang'
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.paragraphs[0].paragraph_format.space_after = Pt(10)  # Adjust the space after each paragraph for bottom padding
                cell.paragraphs[0].paragraph_format.space_before = Pt(10)# Adjust the space after each paragraph for padding
            

    # Prepare the HTTP response with the Word document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=monthly_rental_fee_report_{selected_month}_{selected_year}.docx'
    document.save(response)

    return response



@login_required
def select_month_year(request):
    if request.method == 'POST':
        form = MonthYearForm(request.POST)
        if form.is_valid():
            selected_month = form.cleaned_data['selected_month']
            selected_year = form.cleaned_data['selected_year']

            return redirect('dashboard:monthly_fee_pdf', month=selected_month, year=selected_year)
    else:
        initial_month = datetime.now().month
        initial_year = datetime.now().year
        form = MonthYearForm(initial={'selected_month': initial_month, 'selected_year': initial_year})

    return render(request, 'dashboard/monthly/monthly_rental_fee.html', {'form': form})

def room_pie_chart(request):
    available_count = room.objects.filter(status=room.AVAILABLE).count()
    in_use_count = room.objects.filter(status=room.IN_USE).count()

    data = {
        'labels': ['Available', 'In Use'],
        'data': [available_count, in_use_count],
    }

    return JsonResponse(data)


from django.db.models import Count
from django.db.models.functions import ExtractMonth

def checkin_checkout_bar_chart(request):
    current_month = timezone.now().month
    current_year = timezone.now().year

    checkin_data = CheckIn.objects.filter(date__year=current_year, date__month=current_month).count()
    checkout_data = CheckOut.objects.filter(date__year=current_year, date__month=current_month).count()

    data = {
        'labels': [f'{current_month} {current_year}'],
        'checkin': [{'count': checkin_data}],
        'checkout': [{'count': checkout_data}],
    }

    return JsonResponse(data, safe=False)

def room_checkin_chart(request):
    current_year = datetime.now().year

    # Get all rooms
    rooms = room.objects.all()

    # Count check-ins for each room
    checkin_data = []
    for room_instance in rooms:
        checkin_count = CheckIn.objects.filter(room=room_instance, date__year=current_year).count()
        checkin_data.append({
            'room_no': room_instance.RoomNo,
            'checkin_count': checkin_count,
        })

    data = {
        'labels': [entry['room_no'] for entry in checkin_data],
        'checkin_count': [entry['checkin_count'] for entry in checkin_data],
    }

    return JsonResponse(data, safe=False)
